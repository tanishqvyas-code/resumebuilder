from __future__ import annotations

from io import BytesIO

from app.models.resume import ResumeData


def resume_to_docx(data: ResumeData) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError("python-docx is not installed. Enable it in requirements to export DOCX.") from e

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run(data.personal.full_name.strip() or "Your Name")
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    contact_bits = [
        data.personal.email,
        data.personal.phone,
        data.personal.location,
        data.personal.linkedin_url,
        data.personal.portfolio_url,
    ]
    contact_line = " | ".join(x.strip() for x in contact_bits if x and str(x).strip())
    if contact_line:
        cp = doc.add_paragraph(contact_line)
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_heading(text: str) -> None:
        h = doc.add_paragraph()
        hr = h.add_run(text.upper())
        hr.bold = True
        hr.font.size = Pt(12)

    if data.professional_summary.strip():
        add_heading("Professional Summary")
        for para in data.professional_summary.strip().split("\n\n"):
            doc.add_paragraph(para.strip())

    if data.skills:
        add_heading("Skills")
        for cat in data.skills:
            if not cat.items:
                continue
            line = f"{cat.category_name}: {', '.join(cat.items)}"
            doc.add_paragraph(line)

    if data.work_experience:
        add_heading("Work Experience")
        for job in data.work_experience:
            head = " — ".join(
                x for x in (job.job_title.strip(), job.company_name.strip()) if x
            )
            if head:
                jp = doc.add_paragraph()
                jr = jp.add_run(head)
                jr.bold = True
            meta = []
            if job.location.strip():
                meta.append(job.location.strip())
            dates = " – ".join(x for x in (job.start_date.strip(), job.end_date.strip()) if x)
            if dates:
                meta.append(dates)
            if meta:
                doc.add_paragraph(" | ".join(meta))
            for b in job.bullets:
                if b.strip():
                    doc.add_paragraph(b.strip(), style="List Bullet")

    if data.education:
        add_heading("Education")
        for ed in data.education:
            line = " — ".join(
                x
                for x in (
                    ed.degree.strip(),
                    ed.institution.strip(),
                    ed.location.strip(),
                    ed.graduation_year.strip(),
                )
                if x
            )
            if ed.gpa and str(ed.gpa).strip():
                line = f"{line} (GPA: {ed.gpa})" if line else f"GPA: {ed.gpa}"
            if line:
                doc.add_paragraph(line)

    if data.projects:
        add_heading("Projects")
        for pr in data.projects:
            pp = doc.add_paragraph()
            prun = pp.add_run(pr.title.strip() or "Project")
            prun.bold = True
            if pr.technologies_used.strip():
                doc.add_paragraph(f"Technologies: {pr.technologies_used.strip()}")
            if pr.description.strip():
                doc.add_paragraph(pr.description.strip())
            if pr.link.strip():
                doc.add_paragraph(pr.link.strip())

    if data.certifications:
        add_heading("Certifications")
        for c in data.certifications:
            parts = [c.name.strip(), c.issuer.strip(), c.date.strip()]
            line = " — ".join(x for x in parts if x)
            if line:
                doc.add_paragraph(line)

    if data.achievements:
        add_heading("Achievements")
        for a in data.achievements:
            if a.strip():
                doc.add_paragraph(a.strip(), style="List Bullet")

    if data.languages:
        add_heading("Languages")
        for lang in data.languages:
            line = " — ".join(
                x for x in (lang.language.strip(), lang.proficiency.strip()) if x
            )
            if line:
                doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
